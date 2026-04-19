"""
File Extractor Service  v2  —  Fixed & Enhanced
-----------------------------------------------
Robust text extraction from:
  ✓ PDF files (via PyMuPDF)
  ✓ Word documents (.docx)
  ✓ Images (PNG, JPG, JPEG via OCR)
  ✓ Text files (UTF-8, Latin-1 fallback)

Features:
  - Graceful degradation if libraries missing
  - Whitespace normalization
  - Character count reporting
  - Clear error messages for debugging

Install dependencies:
  pip install pymupdf python-docx pillow pytesseract
For OCR to work, also install Tesseract:
  macOS:    brew install tesseract
  Ubuntu:   sudo apt-get install tesseract-ocr
  Windows:  Download from https://github.com/UB-Mannheim/tesseract/wiki
"""

import io
import os

# Optional imports (graceful fallback)
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[FileExtractor] PyMuPDF not installed. PDF support disabled. Install: pip install pymupdf")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[FileExtractor] python-docx not installed. Word support disabled. Install: pip install python-docx")

try:
    from PIL import Image
    import pytesseract
    # Configure Tesseract path for Windows (common install location)
    import platform
    if platform.system() == "Windows":
        pytesseract.pytesseract.pytesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("[FileExtractor] PIL/pytesseract not installed. OCR support disabled. Install: pip install pillow pytesseract")


# ── Supported formats ───────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg"}


def is_supported(filename: str) -> bool:
    """Check if file type is supported."""
    ext = os.path.splitext(filename.lower())[1]
    return ext in SUPPORTED_EXTENSIONS


# ── Main entry point ───────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> dict:
    """
    Extract text from uploaded file.
    
    Args:
        file_bytes: Raw file bytes
        filename:   Filename (for extension detection)
    
    Returns:
        {
            "success": True/False,
            "text": extracted text (or empty string if failed),
            "method": extraction method used,
            "char_count": number of characters extracted,
            "error": error message (if failed)
        }
    """
    ext = os.path.splitext(filename.lower())[1]
    
    try:
        if ext == ".pdf":
            return _extract_pdf(file_bytes)
        
        elif ext == ".docx":
            return _extract_docx(file_bytes)
        
        elif ext == ".txt":
            return _extract_txt(file_bytes)
        
        elif ext in {".png", ".jpg", ".jpeg"}:
            return _extract_image(file_bytes)
        
        else:
            return _fail(f"Unsupported file extension: {ext}")
    
    except Exception as e:
        return _fail(f"Extraction error: {str(e)}")


# ── PDF extraction ────────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes) -> dict:
    """Extract text from PDF using PyMuPDF."""
    if not PDF_AVAILABLE:
        return _fail(
            "PDF extraction not available. Install PyMuPDF: pip install pymupdf"
        )
    
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        
        # Extract from each page
        for page_num, page in enumerate(doc):
            try:
                page_text = page.get_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                print(f"[PDFExtractor] Warning: Could not extract page {page_num + 1}: {e}")
                continue
        
        text = "\n".join(text_parts)
        doc.close()
        
        return _success(text, method="pdf")
    
    except Exception as e:
        return _fail(f"PDF parsing failed: {str(e)}")


# ── DOCX extraction ──────────────────────────────────────────────────────

def _extract_docx(file_bytes: bytes) -> dict:
    """Extract text from Word document."""
    if not DOCX_AVAILABLE:
        return _fail(
            "Word document extraction not available. Install: pip install python-docx"
        )
    
    try:
        file_stream = io.BytesIO(file_bytes)
        document = docx.Document(file_stream)
        
        # Extract paragraphs
        text_parts = []
        for para in document.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract from tables (if any)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        text = "\n".join(text_parts)
        return _success(text, method="docx")
    
    except Exception as e:
        return _fail(f"Word document parsing failed: {str(e)}")


# ── TXT extraction ────────────────────────────────────────────────────────

def _extract_txt(file_bytes: bytes) -> dict:
    """Extract text from plain text file."""
    try:
        # Try UTF-8 first
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            # Fall back to Latin-1 (always succeeds)
            text = file_bytes.decode("latin-1")
        
        return _success(text, method="txt")
    
    except Exception as e:
        return _fail(f"Text file parsing failed: {str(e)}")


# ── Image OCR ─────────────────────────────────────────────────────────────

def _extract_image(file_bytes: bytes) -> dict:
    """Extract text from image via OCR."""
    if not OCR_AVAILABLE:
        return _fail(
            "Image OCR not available. Install: pip install pillow pytesseract\n"
            "Also install Tesseract binary: "
            "macOS (brew install tesseract), "
            "Ubuntu (sudo apt-get install tesseract-ocr), "
            "Windows (https://github.com/UB-Mannheim/tesseract/releases)"
        )
    
    try:
        # Open image
        image = Image.open(io.BytesIO(file_bytes))
        
        # Convert to RGB if needed (e.g. RGBA, grayscale)
        if image.mode != "RGB":
            image = image.convert("RGB")
        
        # Run OCR
        text = pytesseract.image_to_string(image)
        
        if not text or not text.strip():
            return _fail("Image contains no readable text or Tesseract not found at path. Install: https://github.com/UB-Mannheim/tesseract/releases")
        
        return _success(text, method="ocr")
    
    except FileNotFoundError as e:
        return _fail(f"Tesseract binary not found. Install from: https://github.com/UB-Mannheim/tesseract/releases. Error: {str(e)}")
    except Exception as e:
        return _fail(f"Image OCR failed: {str(e)}")


# ── Helpers ──────────────────────────────────────────────────────────────

def _success(text: str, method: str) -> dict:
    """Return successful extraction result."""
    text = text.strip()
    
    # Normalize whitespace
    text = "\n".join(line.strip() for line in text.split("\n") if line.strip())
    
    return {
        "success": True,
        "text": text,
        "method": method,
        "char_count": len(text),
        "error": None,
    }


def _fail(error: str) -> dict:
    """Return failed extraction result."""
    return {
        "success": False,
        "text": "",
        "method": "none",
        "char_count": 0,
        "error": error,
    }


# ── Self-test ────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("File Extractor v2 — Testing capabilities\n")
    print(f"  PDF support:    {'✓' if PDF_AVAILABLE else '✗'}")
    print(f"  DOCX support:   {'✓' if DOCX_AVAILABLE else '✗'}")
    print(f"  OCR support:    {'✓' if OCR_AVAILABLE else '✗'}")
    print(f"\nSupported formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}\n")
    
    # Test with example content
    examples = {
        "test.txt": b"Patient presents with chest pain and dyspnea",
        "test.pdf": b"",  # Would need real PDF
    }
    
    for filename, content in examples.items():
        if filename.endswith(".txt") and content:
            result = extract_text(content, filename)
            print(f"{filename}:")
            print(f"  Success: {result['success']}")
            print(f"  Method:  {result['method']}")
            print(f"  Chars:   {result['char_count']}")